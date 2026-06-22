#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_chat_threads.py — GSP NEXT 30
Quét file .docx trong data/luong_chat_chu_tich/
Tạo data/luong_chat_chu_tich.json
Nhúng JSON vào index.html giữa marker LCC-DATA-START / LCC-DATA-END

Chạy: python scripts/build_chat_threads.py
"""

import os, json, re, sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx chưa được cài.")
    print("       Chạy: pip install python-docx")
    sys.exit(1)

# ─── Paths ────────────────────────────────────────────────────────────────────
BASE_DIR    = Path(__file__).parent.parent
INPUT_DIR   = BASE_DIR / "data" / "luong_chat_chu_tich"
OUTPUT_JSON = BASE_DIR / "data" / "luong_chat_chu_tich.json"
INDEX_HTML  = BASE_DIR / "index.html"
HTML_MARK_S = "<!-- LCC-DATA-START -->"
HTML_MARK_E = "<!-- LCC-DATA-END -->"

# ─── Từ khóa nhận diện ────────────────────────────────────────────────────────
ACTION_KW = [
    "kế hoạch hành động", "action plan", "công việc cần làm",
    "việc tiếp theo", "việc cần làm", "owner", "deadline",
    "d7", "d14", "d30", "d60", "d90",
    "ngày hoàn thành", "phụ trách:", "hạn:",
    "mốc", "giao cho", "giao việc",
]
CHAIR_KW = [
    "chủ tịch chốt", "chủ tịch đã chốt", "chủ tịch quyết",
    "chủ tịch đồng ý", "chủ tịch phê duyệt", "đã được chốt",
    "đã thống nhất", "kết luận:", "quyết định:", "chốt:",
]
FOLLOW_KW = [
    "cần theo dõi", "cần tiếp tục", "theo dõi thêm",
    "chờ phản hồi", "đang chờ", "cần bổ sung",
    "cập nhật tiếp", "việc tiếp theo",
]
TOPIC_MAP = [
    (["sop", "quy trình", "quy chuẩn", "tiêu chuẩn"],              "Quy trình / SOP"),
    (["kế hoạch", "lộ trình", "roadmap"],                           "Kế hoạch"),
    (["mrp", "aps", "s&op", "lập kế hoạch", "chuỗi cung ứng"],      "Lập kế hoạch / MRP"),
    (["pic", "nhân sự", "nhân lực", "phân công"],                   "Nhân sự / PIC"),
    (["chất lượng", "qms", "ncr", "car", "8d", "fpy", "copq"],      "Chất lượng"),
    (["kho", "vận hành", "nhà máy", "btp", "nvl"],                  "Vận hành / Kho"),
    (["ai", "digital", "claude", "dữ liệu", "ssot", "data"],        "Digital / Dữ liệu / AI"),
    (["pmo", "dự án", "tiến độ", "rag"],                            "PMO"),
    (["kinh doanh", "sales", "khách hàng", "crm", "doanh thu"],     "Kinh doanh / Sales"),
    (["phê duyệt", "quyết định", "chốt"],                           "Quyết định / Phê duyệt"),
]


# ─── Helpers ──────────────────────────────────────────────────────────────────

def file_to_id(filename: str) -> str:
    m = re.match(r'^(\d+)', filename)
    return "CHAT-{:03d}".format(int(m.group(1))) if m else "CHAT-{:03d}".format(abs(hash(filename)) % 900 + 100)


def filename_to_name(filename: str) -> str:
    stem = Path(filename).stem
    stem = re.sub(r'^\d+[_\-]', '', stem)
    return stem.replace('_', ' ').replace('-', ' ')


def get_all_blocks(doc) -> list:
    """Trả về tất cả đoạn văn và ô bảng dưới dạng list dict."""
    blocks = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            blocks.append({'kind': 'para', 'style': para.style.name, 'text': t})
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    blocks.append({'kind': 'cell', 'style': '', 'text': t})
    return blocks


def get_first_heading(doc) -> str:
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        sname = para.style.name.lower()
        if 'heading' in sname or 'title' in sname:
            return t
        # Đoạn đầu to và ngắn (<= 120 ký tự) cũng coi là tiêu đề
        if len(t) <= 120 and (para.runs and any(r.bold for r in para.runs if r.text.strip())):
            return t
    return ''


def detect_topic(blocks: list) -> str:
    haystack = ' '.join(b['text'] for b in blocks[:30]).lower()
    for kws, label in TOPIC_MAP:
        if any(kw in haystack for kw in kws):
            return label
    return 'Chung'


def detect_action_plan(blocks: list) -> tuple:
    """Returns (status: str, summary: str)
       status: 'Có' | 'Không có' | 'Chưa xác định'
    """
    if len(blocks) < 3:
        return 'Chưa xác định', ''

    full_lower = ' '.join(b['text'] for b in blocks).lower()

    if any(kw in full_lower for kw in ACTION_KW):
        # Thu thập dòng có từ khóa action để tóm tắt
        lines = []
        action_kws_short = ['d7','d14','d30','d60','d90','owner','deadline',
                            'phụ trách','hạn:','giao cho','việc tiếp theo','việc cần làm']
        for b in blocks:
            bl = b['text'].lower()
            if any(kw in bl for kw in action_kws_short):
                lines.append(b['text'])
            if len(lines) >= 4:
                break
        summary = '; '.join(lines) if lines else 'Có kế hoạch hành động — xem chi tiết trong file.'
        return 'Có', summary[:400]

    return 'Không có', ''


def detect_chair_confirmed(blocks: list) -> bool:
    full_lower = ' '.join(b['text'] for b in blocks).lower()
    return any(kw in full_lower for kw in CHAIR_KW)


def extract_follow_up(blocks: list) -> str:
    full_lower = ' '.join(b['text'] for b in blocks).lower()
    if not any(kw in full_lower for kw in FOLLOW_KW):
        return ''
    lines = []
    for b in blocks:
        bl = b['text'].lower()
        if any(kw in bl for kw in FOLLOW_KW):
            lines.append(b['text'])
        if len(lines) >= 2:
            break
    return '; '.join(lines)[:300]


def extract_summary(blocks: list) -> str:
    """Lấy 2–4 câu đầu tiên có nội dung thực từ các đoạn (không phải tiêu đề)."""
    candidates = []
    for b in blocks:
        t = b['text']
        # Bỏ qua đoạn quá ngắn hoặc toàn chữ hoa (thường là tiêu đề)
        if len(t) < 25:
            continue
        if t == t.upper() and len(t) < 80:
            continue
        candidates.append(t)
        if len(candidates) >= 3:
            break
    if not candidates:
        candidates = [b['text'] for b in blocks if len(b['text']) > 10][:3]
    return ' '.join(candidates)[:500]


def process_docx(filepath: Path) -> dict:
    filename = filepath.name
    file_id  = file_to_id(filename)
    mod_date = datetime.fromtimestamp(filepath.stat().st_mtime).strftime('%Y-%m-%d')

    doc    = Document(str(filepath))
    blocks = get_all_blocks(doc)

    heading  = get_first_heading(doc)
    ten_luong = heading if heading else filename_to_name(filename)

    chu_de               = detect_topic(blocks)
    tom_tat              = extract_summary(blocks)
    co_ke_hoach, tom_tat_ke = detect_action_plan(blocks)
    co_chot              = detect_chair_confirmed(blocks)
    can_theo_doi         = extract_follow_up(blocks)

    rel_path = str(filepath.relative_to(BASE_DIR)).replace('\\', '/')

    return {
        "id":                    file_id,
        "tenLuongChat":          ten_luong,
        "chuDe":                 chu_de,
        "tomTat":                tom_tat,
        "coKeHoachHanhDong":     co_ke_hoach,
        "tomTatKeHoach":         tom_tat_ke,
        "coNoiDungChuTichChot":  co_chot,
        "noiDungCanTheoDoi":     can_theo_doi,
        "ngayCapNhat":           mod_date,
        "tenFile":               filename,
        "duongDanFile":          rel_path,
        "googleDocsUrl":         ""
    }


def embed_into_html(data: list):
    """Nhúng JSON vào index.html giữa marker LCC-DATA-START/END."""
    if not INDEX_HTML.exists():
        print("  WARN: index.html không tồn tại, bỏ qua nhúng HTML.")
        return

    html = INDEX_HTML.read_text(encoding='utf-8')
    json_str = json.dumps(data, ensure_ascii=False, indent=2)

    replacement = (
        HTML_MARK_S + "\n"
        "  <script type=\"application/json\" id=\"lcc-data\">\n"
        + json_str + "\n"
        "  </script>\n"
        "  " + HTML_MARK_E
    )

    pattern = re.escape(HTML_MARK_S) + r'.*?' + re.escape(HTML_MARK_E)
    if re.search(pattern, html, re.DOTALL):
        new_html = re.sub(pattern, replacement, html, flags=re.DOTALL)
        INDEX_HTML.write_text(new_html, encoding='utf-8')
        print("  HTML: Đã cập nhật dữ liệu nhúng trong index.html")
    else:
        print("  WARN: Không tìm thấy marker LCC-DATA-START/END trong index.html.")
        print("        Chạy lại sau khi marker đã được thêm vào HTML.")


def main():
    divider = "=" * 62
    print(divider)
    print("  build_chat_threads.py — GSP NEXT 30")
    print(f"  Thư mục nguồn: {INPUT_DIR}")
    print(divider)

    if not INPUT_DIR.exists():
        INPUT_DIR.mkdir(parents=True)
        print(f"  Đã tạo thư mục: {INPUT_DIR}")

    # Thu thập file .docx (bỏ qua file tạm ~$...)
    docx_files = sorted([
        f for f in INPUT_DIR.glob("*.docx")
        if not f.name.startswith("~$")
    ])

    print(f"  Tìm thấy: {len(docx_files)} file .docx\n")

    results, errors = [], 0

    for fpath in docx_files:
        try:
            rec = process_docx(fpath)
            results.append(rec)
            plan_icon = "✅" if rec["coKeHoachHanhDong"] == "Có" else ("⚠ " if rec["coKeHoachHanhDong"] == "Chưa xác định" else "○ ")
            chot_icon = " 👑" if rec["coNoiDungChuTichChot"] else ""
            print(f"  OK  {fpath.name}")
            print(f"      {rec['id']} | {rec['tenLuongChat']}")
            print(f"      {plan_icon} {rec['coKeHoachHanhDong']}{chot_icon} | {rec['chuDe']}")
        except Exception as exc:
            errors += 1
            print(f"  ERR {fpath.name}: {exc}")
        print()

    # Ghi JSON
    OUTPUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as fh:
        json.dump(results, fh, ensure_ascii=False, indent=2)
    print(f"  Đã ghi: {OUTPUT_JSON}")

    # Nhúng vào HTML
    embed_into_html(results)

    # Thống kê
    co  = sum(1 for r in results if r["coKeHoachHanhDong"] == "Có")
    khong = sum(1 for r in results if r["coKeHoachHanhDong"] == "Không có")
    chua = sum(1 for r in results if r["coKeHoachHanhDong"] == "Chưa xác định")
    chot = sum(1 for r in results if r["coNoiDungChuTichChot"])

    print()
    print(divider)
    print(f"  File đã đọc:               {len(results)}")
    print(f"  File lỗi:                  {errors}")
    print(f"  Luồng có kế hoạch HD:      {co}")
    print(f"  Luồng không có KH:         {khong}")
    print(f"  Luồng chưa xác định KH:    {chua}")
    print(f"  Luồng Chủ tịch đã chốt:    {chot}")
    print(divider)
    if results:
        print("  Tiếp theo: clasp push --force")
    print()


if __name__ == '__main__':
    main()
